from abc import ABC, abstractmethod
import os
import chardet
from PyPDF2 import PdfReader
from typing import Dict, Any

# Try to import docx, handle gracefully if not available
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class FileParser(ABC):
    """
    Abstract base class for file parsers.
    Follows the Open-Closed Principle - open for extension, closed for modification.
    """
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self._validate_file()
    
    def _validate_file(self) -> None:
        """Validate that the file exists and is readable."""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        if not os.path.isfile(self.file_path):
            raise ValueError(f"Path is not a file: {self.file_path}")
    
    @abstractmethod
    def parse(self) -> str:
        """
        Parse the file and extract text content.
        
        Returns:
            str: Extracted text content
        
        Raises:
            Exception: If parsing fails
        """
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> list[str]:
        """
        Get list of supported file extensions for this parser.
        
        Returns:
            list[str]: List of supported extensions (e.g., ['.txt', '.text'])
        """
        pass
    
    def get_file_info(self) -> Dict[str, Any]:
        """
        Get basic information about the file.
        
        Returns:
            dict: File information including size, extension, etc.
        """
        if not os.path.exists(self.file_path):
            return None
        
        stat = os.stat(self.file_path)
        file_extension = os.path.splitext(self.file_path)[1].lower()
        
        return {
            'filename': os.path.basename(self.file_path),
            'extension': file_extension,
            'size_bytes': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'modified_time': stat.st_mtime
        }


class TxtFileParser(FileParser):
    """
    Concrete parser for plain text files.
    Handles encoding detection and various text formats.
    """
    
    def get_supported_extensions(self) -> list[str]:
        return ['.txt', '.text']
    
    def parse(self) -> str:
        """
        Parse a text file with automatic encoding detection.
        
        Returns:
            str: File content as text
        
        Raises:
            Exception: If text parsing fails
        """
        try:
            # Detect encoding
            with open(self.file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            try:
                with open(self.file_path, 'r', encoding=encoding) as file:
                    content = file.read()
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
            
            return content.strip()
        
        except Exception as e:
            raise Exception(f"TXT parsing error: {str(e)}")


class PdfFileParser(FileParser):
    """
    Concrete parser for PDF files.
    Extracts text from all pages in the document.
    """
    
    def get_supported_extensions(self) -> list[str]:
        return ['.pdf']
    
    def parse(self) -> str:
        """
        Parse a PDF file and extract text content from all pages.
        
        Returns:
            str: Extracted text content
        
        Raises:
            Exception: If PDF parsing fails
        """
        try:
            reader = PdfReader(self.file_path)
            text_content = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    # Warning: Could not extract text from page {page_num + 1}: {str(e)}
                    continue
            
            if not text_content:
                raise Exception("No readable text found in PDF")
            
            return '\n\n'.join(text_content).strip()
        
        except Exception as e:
            raise Exception(f"PDF parsing error: {str(e)}")


class DocxFileParser(FileParser):
    """
    Concrete parser for Microsoft Word DOCX files.
    Extracts text from paragraphs and tables.
    """
    
    def get_supported_extensions(self) -> list[str]:
        return ['.docx']
    
    def parse(self) -> str:
        """
        Parse a DOCX file and extract text content from paragraphs and tables.
        
        Returns:
            str: Extracted text content
        
        Raises:
            Exception: If DOCX parsing fails
        """
        if not DOCX_AVAILABLE:
            raise Exception("DOCX parsing not available. Please install python-docx: pip install python-docx")
        
        try:
            doc = Document(self.file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise Exception("No readable text found in DOCX file")
            
            return '\n\n'.join(text_content).strip()
        
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")


class FileParserFactory:
    """
    Factory class for creating appropriate file parsers.
    Implements the Factory Method pattern.
    """
    
    # Registry of available parsers
    _parsers = {
        '.txt': TxtFileParser,
        '.text': TxtFileParser,
        '.pdf': PdfFileParser,
    }
    
    # Add DocxFileParser only if docx library is available
    if DOCX_AVAILABLE:
        _parsers['.docx'] = DocxFileParser
    
    @classmethod
    def create_parser(cls, file_path: str, file_extension: str = None) -> FileParser:
        """
        Create the appropriate parser for the given file.
        
        Args:
            file_path (str): Path to the file
            file_extension (str, optional): File extension. If not provided, 
                                          will be extracted from file_path
        
        Returns:
            FileParser: Appropriate parser instance
        
        Raises:
            ValueError: If file extension is not supported
        """
        if file_extension is None:
            file_extension = os.path.splitext(file_path)[1].lower()
        else:
            file_extension = file_extension.lower()
        
        if file_extension not in cls._parsers:
            supported_extensions = list(cls._parsers.keys())
            raise ValueError(
                f"Unsupported file extension: {file_extension}. "
                f"Supported extensions: {', '.join(supported_extensions)}"
            )
        
        parser_class = cls._parsers[file_extension]
        return parser_class(file_path)
    
    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        """
        Get all supported file extensions.
        
        Returns:
            list[str]: List of all supported extensions
        """
        return list(cls._parsers.keys())
    
    @classmethod
    def register_parser(cls, extension: str, parser_class: type) -> None:
        """
        Register a new parser for a file extension.
        This allows for easy extension without modifying existing code.
        
        Args:
            extension (str): File extension (e.g., '.pptx')
            parser_class (type): Parser class that inherits from FileParser
        
        Raises:
            TypeError: If parser_class doesn't inherit from FileParser
        """
        if not issubclass(parser_class, FileParser):
            raise TypeError("Parser class must inherit from FileParser")
        
        cls._parsers[extension.lower()] = parser_class
    
    @classmethod
    def is_supported(cls, file_extension: str) -> bool:
        """
        Check if a file extension is supported.
        
        Args:
            file_extension (str): File extension to check
        
        Returns:
            bool: True if supported, False otherwise
        """
        return file_extension.lower() in cls._parsers